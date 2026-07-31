"""Parser dokumentow Microsoft Word w formacie Office Open XML (.docx, .docm).

Tresc jest odczytywana w kolejnosci wystapienia w ciele dokumentu, dzieki czemu
tabele zachowuja powiazanie z poprzedzajacym je naglowkiem i akapitem. Parser nie
renderuje strony, wiec nigdy nie zglasza potrzeby OCR: dokument bez tekstu jest
raportowany jako pusty.
"""

from __future__ import annotations

import datetime as _dt
import re
import zipfile
from collections.abc import Iterator
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document as load_document
from docx.document import Document as WordDocument
from docx.opc.exceptions import OpcError, PackageNotFoundError
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from finddocs.errors import (
    CorruptedFileError,
    EmptyDocumentError,
    ExtractionError,
    FindDocsError,
    PasswordProtectedError,
)
from finddocs.extractors.base import ExtractionContext, Extractor
from finddocs.normalization.text import clean_text, fold_diacritics, looks_like_garbage
from finddocs.types import (
    DocumentMetadata,
    ExtractedSection,
    ExtractionResult,
    SupportLevel,
    TextOrigin,
)

#: Nazwy elementow WordprocessingML uzywane przy przechodzeniu po ciele dokumentu.
_TAG_PARAGRAPH = qn("w:p")
_TAG_TABLE = qn("w:tbl")
_TAG_SDT = qn("w:sdt")
_TAG_SDT_CONTENT = qn("w:sdtContent")
_TAG_DRAWING = qn("w:drawing")
_TAG_PICT = qn("w:pict")
_TAG_TBL_CAPTION = qn("w:tblCaption")
_TAG_TBL_DESCRIPTION = qn("w:tblDescription")
_ATTR_VAL = qn("w:val")

#: Prefiksy nazw stylow oznaczajacych nagłówek, po zlozeniu znakow diakrytycznych.
_HEADING_PREFIXES: tuple[str, ...] = ("heading", "nagl")

#: Sygnatura kontenera OLE. Zaszyfrowany plik .docx jest zapisany wlasnie tak.
_OLE_SIGNATURE = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"

#: Co ile przetworzonych elementow sprawdzac anulowanie i limit czasu.
_CHECKPOINT_EVERY = 16

#: Maksymalna dlugosc komorki, ktora moze byc jeszcze nazwa kolumny.
_MAX_HEADER_CELL_CHARS = 64

#: Maksymalna dlugosc zapamietanego naglowka przypisywanego kolejnym sekcjom.
_MAX_HEADING_CHARS = 200

#: Ile znakow tresci wystarczy do oceny, czy tekst nie jest smieciem.
_GARBAGE_SAMPLE_CHARS = 20_000

#: Nazwa czesci pakietu z rozszerzonymi wlasciwosciami dokumentu.
_APP_PART_NAME = "/docProps/app.xml"

#: Interesujace nas pola z docProps/app.xml.
_APP_KEYS = frozenset({"Pages", "Words", "Characters", "Application", "Company"})

_NUMERIC_NOISE_RE = re.compile(r"[\s.,%+\-()]")


@dataclass(slots=True)
class _Collector:
    """Zbiera sekcje dokumentu z zachowaniem limitow kontekstu ekstrakcji."""

    context: ExtractionContext
    sections: list[ExtractedSection] = field(default_factory=list)
    heading: str | None = None
    chars: int = 0
    truncated: bool = False
    steps: int = 0

    def tick(self) -> None:
        """Okresowo sprawdza anulowanie i limit czasu."""
        self.steps += 1
        if self.steps % _CHECKPOINT_EVERY == 0:
            self.context.checkpoint()

    def add(
        self,
        text: str,
        *,
        kind: str = "text",
        heading: str | None = None,
        row: int | None = None,
        extra: dict[str, Any] | None = None,
    ) -> None:
        """Dodaje sekcje po oczyszczeniu tekstu. Puste sekcje sa pomijane."""
        if self.truncated:
            return
        cleaned = clean_text(text)
        if not cleaned:
            return
        remaining = self.context.max_chars - self.chars
        if remaining <= 0:
            self.truncated = True
            return
        if len(cleaned) > remaining:
            cleaned = cleaned[:remaining].rstrip()
            self.truncated = True
            if not cleaned:
                return
        self.chars += len(cleaned)
        self.sections.append(
            ExtractedSection(
                text=cleaned,
                kind=kind,
                order=len(self.sections),
                row=row,
                origin=TextOrigin.NATIVE,
                heading=heading,
                extra=extra or {},
            )
        )

    def set_heading(self, text: str) -> None:
        """Zapamietuje biezacy nagłówek dziedziczony przez kolejne sekcje."""
        cleaned = clean_text(text)
        if cleaned:
            self.heading = cleaned[:_MAX_HEADING_CHARS]


class DocxExtractor(Extractor):
    """Adapter dokumentow Word oparty o biblioteke python-docx."""

    name = "docx"
    extensions = (".docx", ".docm")
    mime_types = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/vnd.ms-word.document.macroEnabled.12",
    )
    support_level = SupportLevel.FULL
    priority = 120

    def extract(self, path: Path, context: ExtractionContext) -> ExtractionResult:
        """Wyciaga tekst, tabele, naglowki stron i metadane z pliku Word."""
        context.checkpoint()
        document = self._open(path)
        collector = _Collector(context=context)
        try:
            self._read_body(document, collector)
            self._read_headers_and_footers(document, collector)
        except FindDocsError:
            raise
        except (
            OpcError,
            zipfile.BadZipFile,
            KeyError,
            ValueError,
            IndexError,
            AttributeError,
        ) as exc:
            raise CorruptedFileError(
                "Struktura dokumentu Word jest uszkodzona, nie udało się odczytać treści.",
                details={"plik": path.name},
            ) from exc
        except OSError as exc:
            raise ExtractionError(
                "Nie udało się odczytać zawartości dokumentu Word.",
                details={"plik": path.name},
            ) from exc
        except Exception as exc:
            raise ExtractionError(
                "Nieoczekiwany błąd podczas odczytu dokumentu Word.",
                details={"plik": path.name, "typ_bledu": type(exc).__name__},
            ) from exc

        if not collector.sections:
            raise self._empty_error(document, path)

        result = ExtractionResult(
            sections=collector.sections,
            origin=TextOrigin.NATIVE,
            parser_name=self.name,
            support_level=self.support_level,
        )
        app_properties = _read_app_properties(document)
        result.metadata = self._read_metadata(document, app_properties, result)
        pages = _page_count(app_properties)
        if pages is not None:
            result.total_pages = pages
            result.metadata.page_count = pages
        if collector.truncated:
            result.warnings.append(
                "Dokument przekroczył limit długości tekstu, zaindeksowano tylko początek."
            )
        sample = _sample_text(collector.sections)
        if sample and looks_like_garbage(sample):
            result.warnings.append(
                "Tekst dokumentu wygląda na uszkodzony, udział znaków alfanumerycznych jest niski."
            )
        return result

    # --- otwieranie pliku ------------------------------------------------------

    def _open(self, path: Path) -> WordDocument:
        """Otwiera pakiet Word i tlumaczy bledy biblioteki na wyjatki FindDocs."""
        if _is_ole_container(path):
            raise PasswordProtectedError(
                "Dokument Word jest zaszyfrowany lub zabezpieczony hasłem.",
                details={"plik": path.name},
            )
        try:
            return load_document(str(path))
        except PackageNotFoundError as exc:
            raise CorruptedFileError(
                "Plik nie jest poprawnym pakietem Word albo jest uszkodzony.",
                details={"plik": path.name},
            ) from exc
        except (zipfile.BadZipFile, KeyError) as exc:
            raise CorruptedFileError(
                "Archiwum dokumentu Word jest niekompletne, brakuje wymaganych części.",
                details={"plik": path.name},
            ) from exc
        except (OpcError, AttributeError, IndexError) as exc:
            raise CorruptedFileError(
                "Struktura pakietu Word jest nieprawidłowa.",
                details={"plik": path.name},
            ) from exc
        except ValueError as exc:
            raise ExtractionError(
                "Nie udało się otworzyć dokumentu Word.",
                details={"plik": path.name},
            ) from exc
        except OSError as exc:
            raise ExtractionError(
                "Nie udało się odczytać pliku dokumentu Word.",
                details={"plik": path.name},
            ) from exc
        except Exception as exc:
            raise ExtractionError(
                "Nieoczekiwany błąd podczas otwierania dokumentu Word.",
                details={"plik": path.name, "typ_bledu": type(exc).__name__},
            ) from exc

    def _empty_error(self, document: WordDocument, path: Path) -> EmptyDocumentError:
        """Buduje wyjatek dla dokumentu bez tresci tekstowej."""
        images = _count_images(document)
        if images:
            message = (
                f"Dokument Word nie zawiera tekstu, znaleziono {images} obiektow graficznych. "
                "Tresc jest prawdopodobnie wstawiona jako obraz."
            )
        else:
            message = "Dokument Word nie zawiera żadnej treści tekstowej."
        return EmptyDocumentError(
            message,
            details={"plik": path.name, "obrazy": images, "parser": self.name},
        )

    # --- tresc -----------------------------------------------------------------

    def _read_body(self, document: WordDocument, collector: _Collector) -> None:
        """Przechodzi po ciele dokumentu, zachowujac kolejnosc akapitow i tabel."""
        for block in _iter_blocks(document.element.body, document):
            collector.tick()
            if collector.truncated:
                return
            if isinstance(block, Paragraph):
                self._add_paragraph(block, collector)
            else:
                self._add_table(block, collector)

    def _add_paragraph(self, paragraph: Paragraph, collector: _Collector) -> None:
        """Dodaje akapit i aktualizuje biezacy nagłówek, gdy akapit jest naglowkiem."""
        text = paragraph.text
        if not text.strip():
            return
        style = _heading_style(paragraph)
        if style is not None:
            collector.set_heading(text)
            collector.add(text, heading=collector.heading, extra={"styl": style})
            return
        collector.add(text, heading=collector.heading)

    def _add_table(self, table: Table, collector: _Collector) -> None:
        """Zamienia tabele na sekcje wierszy, w miare mozliwosci z nazwami kolumn."""
        rows = _table_rows(table, collector)
        if not rows:
            return
        caption = _table_caption(table) or collector.heading
        header = rows[0] if _looks_like_header(rows) else None
        if header is not None:
            collector.add(
                " | ".join(cell for cell in header if cell),
                kind="table_header",
                heading=caption,
                row=1,
            )
        data_rows = rows[1:] if header is not None else rows
        first_number = 2 if header is not None else 1
        for offset, cells in enumerate(data_rows):
            collector.tick()
            if collector.truncated:
                return
            collector.add(
                _format_row(header, cells),
                kind="table_row",
                heading=caption,
                row=first_number + offset,
            )

    def _read_headers_and_footers(self, document: WordDocument, collector: _Collector) -> None:
        """Dopisuje naglowki i stopki sekcji dokumentu, pomijajac powtorzenia."""
        seen: set[str] = set()
        for index, section in enumerate(document.sections, start=1):
            collector.tick()
            if collector.truncated:
                return
            for label, container in (("nagłówek", section.header), ("stopka", section.footer)):
                text = _container_text(container)
                if not text:
                    continue
                key = f"{label}:{text.casefold()}"
                if key in seen:
                    continue
                seen.add(key)
                collector.add(
                    text,
                    extra={"zrodlo": label, "sekcja_dokumentu": index},
                )

    # --- metadane --------------------------------------------------------------

    def _read_metadata(
        self,
        document: WordDocument,
        app_properties: dict[str, str],
        result: ExtractionResult,
    ) -> DocumentMetadata:
        """Czyta wlasciwosci pakietu. Braki nie sa bledem, tylko ostrzezeniem."""
        metadata = DocumentMetadata()
        application = app_properties.get("Application")
        if application:
            metadata.producer = application
        company = app_properties.get("Company")
        if company:
            metadata.extra["firma"] = company
        try:
            properties = document.core_properties
            metadata.title = _text_or_none(properties.title)
            metadata.author = _text_or_none(properties.author)
            metadata.subject = _text_or_none(properties.subject)
            metadata.keywords = _text_or_none(properties.keywords)
            metadata.language = _text_or_none(properties.language)
            metadata.created_at = _as_datetime(properties.created)
            metadata.modified_at = _as_datetime(properties.modified)
            editor = _text_or_none(properties.last_modified_by)
            if editor:
                metadata.extra["ostatnio_zapisal"] = editor
            category = _text_or_none(properties.category)
            if category:
                metadata.extra["kategoria"] = category
        except (OpcError, KeyError, ValueError, AttributeError):
            result.warnings.append("Nie udało się odczytać właściwości dokumentu Word.")
        return metadata


# --- przechodzenie po strukturze dokumentu -------------------------------------


def _iter_blocks(element: Any, parent: Any) -> Iterator[Paragraph | Table]:
    """Zwraca akapity i tabele w kolejnosci wystapienia w podanym elemencie.

    Zaglada do kontrolek zawartosci (``w:sdt``), bo Word chowa w nich zwykle
    akapity i tabele formularzy.
    """
    for child in element:
        tag = child.tag
        if tag == _TAG_PARAGRAPH:
            yield Paragraph(child, parent)
        elif tag == _TAG_TABLE:
            yield Table(child, parent)
        elif tag == _TAG_SDT:
            content = child.find(_TAG_SDT_CONTENT)
            if content is not None:
                yield from _iter_blocks(content, parent)


def _heading_style(paragraph: Paragraph) -> str | None:
    """Zwraca nazwe stylu, gdy akapit jest naglowkiem, w przeciwnym razie None."""
    try:
        style = paragraph.style
    except (KeyError, ValueError, AttributeError):
        return None
    if style is None:
        return None
    name = str(style.name or "")
    identifier = str(getattr(style, "style_id", "") or "")
    for candidate in (name, identifier):
        if not candidate:
            continue
        folded = fold_diacritics(candidate).casefold().lstrip()
        if folded.startswith(_HEADING_PREFIXES):
            return name or identifier
    return None


def _table_rows(table: Table, collector: _Collector) -> list[list[str]]:
    """Zwraca tresc tabeli jako liste wierszy. Scalone komorki nie sa powielane."""
    rows: list[list[str]] = []
    try:
        table_rows = list(table.rows)
    except (ValueError, IndexError, AttributeError):
        return rows
    for row in table_rows:
        collector.tick()
        if collector.truncated:
            break
        try:
            cells = list(row.cells)
        except (ValueError, IndexError, AttributeError):
            continue
        values: list[str] = []
        previous: Any = None
        for cell in cells:
            element = getattr(cell, "_tc", None)
            if element is not None and element is previous:
                continue
            previous = element
            values.append(_flatten(cell.text))
        rows.append(values)
    return rows


def _table_caption(table: Table) -> str | None:
    """Nazwa tabeli zapisana we wlasciwosciach dostepnosci, gdy autor ja podal."""
    properties = getattr(table._tbl, "tblPr", None)
    if properties is None:
        return None
    for tag in (_TAG_TBL_CAPTION, _TAG_TBL_DESCRIPTION):
        node = properties.find(tag)
        if node is None:
            continue
        value = node.get(_ATTR_VAL)
        if not value:
            continue
        caption = clean_text(str(value))
        if caption:
            return caption[:_MAX_HEADING_CHARS]
    return None


def _looks_like_header(rows: list[list[str]]) -> bool:
    """Ocenia, czy pierwszy wiersz tabeli jest wierszem nazw kolumn."""
    if len(rows) < 2:
        return False
    first = rows[0]
    filled = [cell for cell in first if cell]
    if not filled:
        return False
    if len(filled) * 2 < len(first):
        return False
    if any(len(cell) > _MAX_HEADER_CELL_CHARS for cell in filled):
        return False
    if not any(any(ch.isalpha() for ch in cell) for cell in filled):
        return False
    return not any(_is_numeric(cell) for cell in filled)


def _is_numeric(value: str) -> bool:
    """True dla komorek zawierajacych wylacznie liczbe, kwote albo procent."""
    digits = _NUMERIC_NOISE_RE.sub("", value)
    return bool(digits) and digits.isdigit()


def _format_row(header: list[str] | None, cells: list[str]) -> str:
    """Sklada wiersz tabeli w linie tekstu, z nazwami kolumn gdy sa dostepne."""
    parts: list[str] = []
    for index, value in enumerate(cells):
        if not value:
            continue
        label = header[index] if header is not None and index < len(header) else ""
        parts.append(f"{label}: {value}" if label else value)
    return " | ".join(parts)


def _container_text(container: Any) -> str:
    """Tekst naglowka albo stopki: akapity i wiersze tabel w jednym bloku."""
    parts: list[str] = []
    try:
        paragraphs = list(container.paragraphs)
        tables = list(container.tables)
    except (ValueError, IndexError, AttributeError):
        return ""
    for paragraph in paragraphs:
        value = _flatten(paragraph.text)
        if value:
            parts.append(value)
    for table in tables:
        try:
            table_rows = list(table.rows)
        except (ValueError, IndexError, AttributeError):
            continue
        for row in table_rows:
            try:
                cells = list(row.cells)
            except (ValueError, IndexError, AttributeError):
                continue
            line = " | ".join(value for value in (_flatten(c.text) for c in cells) if value)
            if line:
                parts.append(line)
    return "\n".join(parts)


def _count_images(document: WordDocument) -> int:
    """Liczy obiekty graficzne w ciele dokumentu, uzywane w komunikacie o pustce."""
    body = document.element.body
    total = 0
    for tag in (_TAG_DRAWING, _TAG_PICT):
        total += sum(1 for _ in body.iter(tag))
    return total


def _read_app_properties(document: WordDocument) -> dict[str, str]:
    """Czyta docProps/app.xml. Brak tej czesci nie jest bledem."""
    values: dict[str, str] = {}
    try:
        for part in document.part.package.iter_parts():
            if str(part.partname) != _APP_PART_NAME:
                continue
            root = parse_xml(part.blob)
            for child in root:
                tag = str(child.tag)
                local = tag.rsplit("}", 1)[-1]
                text = child.text
                if local in _APP_KEYS and text and text.strip():
                    values[local] = text.strip()
            break
    except (OpcError, KeyError, ValueError, SyntaxError, AttributeError, OSError):
        return {}
    return values


def _page_count(app_properties: dict[str, str]) -> int | None:
    """Liczba stron zapamietana przez Worda przy ostatnim zapisie."""
    raw = app_properties.get("Pages", "")
    if raw.isdigit():
        pages = int(raw)
        if pages > 0:
            return pages
    return None


# --- drobne narzedzia -----------------------------------------------------------


def _is_ole_container(path: Path) -> bool:
    """Zaszyfrowany dokument OOXML jest zapisany jako kontener OLE, nie jako ZIP."""
    try:
        with path.open("rb") as handle:
            return handle.read(len(_OLE_SIGNATURE)) == _OLE_SIGNATURE
    except OSError:
        return False


def _flatten(text: str) -> str:
    """Skleja tekst w jedna linie, uzywane dla komorek tabeli."""
    return " ".join(text.split())


def _text_or_none(value: str | None) -> str | None:
    """Oczyszczona wartosc metadanej albo None, gdy pole jest puste."""
    if not value:
        return None
    cleaned = _flatten(clean_text(value))
    return cleaned or None


def _as_datetime(value: _dt.datetime | None) -> _dt.datetime | None:
    """Przepuszcza tylko poprawne daty, bo wlasciwosci pakietu bywaja uszkodzone."""
    if isinstance(value, _dt.datetime):
        return value
    return None


def _sample_text(sections: list[ExtractedSection], limit: int = _GARBAGE_SAMPLE_CHARS) -> str:
    """Poczatkowy fragment tresci uzywany do oceny jakosci tekstu."""
    parts: list[str] = []
    total = 0
    for section in sections:
        parts.append(section.text)
        total += len(section.text)
        if total >= limit:
            break
    return " ".join(parts)


__all__ = ["DocxExtractor"]
