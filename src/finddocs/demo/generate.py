"""Generator polskiego zbioru demonstracyjnego FindDocs.

Zbior pozwala uruchomic aplikacje bez dostepu do SharePointa i przetestowac
indeksowanie, OCR, trzy tryby wyszukiwania oraz raport pokrycia. Powstaje
struktura katalogow z dokumentami w formatach DOCX, TXT, RTF, PDF, CSV, XLSX,
EML, PNG i HTML, uzupelniona o pliki celowo problematyczne.

Wszystkie dane sa calkowicie fikcyjne. Nazwiska, firmy, numery rachunkow
i identyfikatory zostaly wymyslone na potrzeby testow, numery zaczynaja sie od
cyfr 00 albo 99, a kazdy dokument zawiera zdanie informujace o tym wprost.

Modul nie korzysta z sieci. Pliki PDF sa skladane recznie, bo biblioteka
pypdfium2 sluzy wylacznie do odczytu: wersja tekstowa uzywa czcionki Helvetica
z wlasnym kodowaniem i mapa ToUnicode dla polskich znakow, a wersja udajaca skan
osadza obraz JPEG bez warstwy tekstowej.
"""

from __future__ import annotations

import csv
import datetime as _dt
import io
import json
import random
import shutil
import struct
import textwrap
import zlib
from dataclasses import dataclass, field
from email.message import EmailMessage
from email.utils import format_datetime
from pathlib import Path
from typing import Any, Final

from finddocs.app_paths import AppPaths
from finddocs.errors import (
    ConfigurationError,
    DependencyUnavailableError,
    TemporaryStorageError,
)
from finddocs.logging_setup import get_logger
from finddocs.version import APP_VERSION

log = get_logger(__name__)

# --- stale zbioru ---------------------------------------------------------------

#: Nazwa pliku z opisem zbioru. Nie jest liczona jako dokument zbioru.
MANIFEST_NAME: Final = "manifest.json"

#: Wersja formatu manifestu.
MANIFEST_VERSION: Final = 1

#: Zdanie powtarzane w kazdym dokumencie zbioru.
DISCLAIMER: Final = "Dokument testowy, dane fikcyjne."

#: Testowy numer rachunku w trzech wariantach zapisu.
ACCOUNT_COMPACT: Final = "00123456789012345678901234"
ACCOUNT_SPACED: Final = "00 1234 5678 9012 3456 7890 1234"
ACCOUNT_DASHED: Final = "00-1234-5678-9012-3456-7890-1234"
ACCOUNT_VARIANTS: Final[tuple[str, ...]] = (ACCOUNT_SPACED, ACCOUNT_DASHED, ACCOUNT_COMPACT)

#: Liczba dokumentow zawierajacych testowy numer rachunku. Punkt odniesienia testow.
EXPECTED_ACCOUNT_DOCUMENTS: Final = 17

#: Rachunki pomocnicze. Nigdy nie zawieraja ciagu cyfr rachunku testowego.
ACCOUNT_B_SPACED: Final = "99 8765 4321 0987 6543 2109 8765"
ACCOUNT_TECH_SPACED: Final = "99 4321 8765 5555 6666 7777 8888"

#: Fragment numeru karty uzywany w zapytaniach o platnosc kartowa.
CARD_SUFFIX: Final = "384675"
CARD_MASKED: Final = "****384675"

#: Fikcyjne podmioty i osoby.
CLIENT_A: Final = "ACME Polska sp. z o.o."
CLIENT_B: Final = "Nowak-Bud sp. j."
BANK: Final = "Bank Testowy S.A."
PERSON_A: Final = "Kowalski Jan"
PERSON_B: Final = "Nowak Anna"
PERSON_C: Final = "Wiśniewska Katarzyna"
PERSON_D: Final = "Zieliński Marek"

#: Numery umow uzywane w zapytaniach dokladnych.
CONTRACT_ACCOUNT: Final = "00/RB/2005/117"
CONTRACT_CREDIT: Final = "99/KR/2015/042"
INVOICE_ID: Final = "FV/00/2015/07"

# --- sciezki wzgledne dokumentow ------------------------------------------------

DOC_PROC_PRZELEWY: Final = "procedury/procedura-przelewow-2015.docx"
DOC_PROC_AML: Final = "procedury/procedura-przeciwdzialania-praniu-pieniedzy.docx"
DOC_PROC_REKLAMACJE: Final = "procedury/procedura-obslugi-reklamacji.docx"
DOC_PROC_KSIEGOWANIE: Final = "procedury/instrukcja-ksiegowania-2007.txt"
DOC_PROC_ARCHIWIZACJA: Final = "procedury/procedura-archiwizacji.txt"
DOC_UMOWA_RACHUNKU: Final = "umowy/umowa-rachunku-acme.docx"
DOC_UMOWA_KREDYTU: Final = "umowy/umowa-kredytowa-nowak-bud.docx"
DOC_ANEKS: Final = "umowy/aneks-nr-1-do-umowy-rachunku.rtf"
DOC_POROZUMIENIE: Final = "umowy/porozumienie-rozliczeniowe.pdf"
DOC_TRX_A_2015: Final = "transakcje/klientA/transakcje-2015-07.csv"
DOC_TRX_A_2007: Final = "transakcje/klientA/transakcje-2007-05.csv"
DOC_TRX_B_2015: Final = "transakcje/klientB/transakcje-2015-07.csv"
DOC_TRX_B_XLSX: Final = "transakcje/klientB/zestawienie-klientB.xlsx"
DOC_MAIL_POTWIERDZENIE: Final = "korespondencja/potwierdzenie-przelewu.eml"
DOC_MAIL_ZAPYTANIE: Final = "korespondencja/zapytanie-o-ksiegowania.eml"
DOC_SKAN_PNG: Final = "skany/skan-potwierdzenia-wplaty.png"
DOC_SKAN_PDF: Final = "skany/skan-umowy-o-wspolpracy.pdf"
DOC_RAPORT_XLSX: Final = "raporty/raport-transakcji-2015.xlsx"
DOC_RAPORT_CSV: Final = "raporty/zestawienie-powiazan.csv"
DOC_RAPORT_HTML: Final = "raporty/podsumowanie-kwartalne.html"
DOC_BLAD_PDF: Final = "problemy/uszkodzony-raport.pdf"
DOC_BLAD_TXT: Final = "problemy/pusty-dokument.txt"
DOC_BLAD_XYZ: Final = "problemy/dane-eksportu.xyz"
DOC_BLAD_DOCX: Final = "problemy/zabezpieczona-umowa.docx"

#: Dokumenty zawierajace testowy numer rachunku, w podziale na warianty zapisu.
ACCOUNT_DOCUMENTS_SPACED: Final[tuple[str, ...]] = (
    DOC_PROC_PRZELEWY,
    DOC_PROC_ARCHIWIZACJA,
    DOC_UMOWA_RACHUNKU,
    DOC_POROZUMIENIE,
    DOC_TRX_A_2015,
    DOC_TRX_B_2015,
    DOC_MAIL_POTWIERDZENIE,
    DOC_RAPORT_XLSX,
    DOC_RAPORT_CSV,
)
ACCOUNT_DOCUMENTS_DASHED: Final[tuple[str, ...]] = (
    DOC_PROC_AML,
    DOC_UMOWA_KREDYTU,
    DOC_ANEKS,
    DOC_MAIL_ZAPYTANIE,
)
ACCOUNT_DOCUMENTS_COMPACT: Final[tuple[str, ...]] = (
    DOC_PROC_REKLAMACJE,
    DOC_PROC_KSIEGOWANIE,
    DOC_TRX_A_2007,
    DOC_TRX_B_XLSX,
)
ACCOUNT_DOCUMENTS: Final[tuple[str, ...]] = (
    ACCOUNT_DOCUMENTS_SPACED + ACCOUNT_DOCUMENTS_DASHED + ACCOUNT_DOCUMENTS_COMPACT
)

#: Dokumenty tekstowe zawierajace zdanie o fikcyjnosci danych, bez skanow i plikow wadliwych.
DISCLAIMER_DOCUMENTS: Final[tuple[str, ...]] = (
    DOC_PROC_PRZELEWY,
    DOC_PROC_AML,
    DOC_PROC_REKLAMACJE,
    DOC_PROC_KSIEGOWANIE,
    DOC_PROC_ARCHIWIZACJA,
    DOC_UMOWA_RACHUNKU,
    DOC_UMOWA_KREDYTU,
    DOC_ANEKS,
    DOC_POROZUMIENIE,
    DOC_TRX_A_2015,
    DOC_TRX_A_2007,
    DOC_TRX_B_2015,
    DOC_TRX_B_XLSX,
    DOC_MAIL_POTWIERDZENIE,
    DOC_MAIL_ZAPYTANIE,
    DOC_RAPORT_XLSX,
    DOC_RAPORT_CSV,
    DOC_RAPORT_HTML,
)


# --- struktury opisu zbioru -----------------------------------------------------


@dataclass(slots=True)
class DemoQuery:
    """Zapytanie referencyjne wraz z lista dokumentow istotnych."""

    query: str
    """Tresc zapytania w postaci wpisywanej przez uzytkownika."""

    mode: str
    """Tryb wyszukiwania: exact, semantic albo hybrid."""

    expected_paths: list[str]
    """Sciezki wzgledne dokumentow, ktore powinny znalezc sie w wynikach."""

    description: str
    """Co dane zapytanie sprawdza. Opisy zaczynajace sie od 'Wymaga OCR' dotycza skanow."""


@dataclass(slots=True)
class DemoCorpusInfo:
    """Opis wygenerowanego zbioru demonstracyjnego."""

    root: Path
    """Katalog glowny zbioru."""

    files: int
    """Liczba wygenerowanych dokumentow. Plik manifestu nie jest liczony."""

    by_extension: dict[str, int]
    """Liczba plikow wedlug rozszerzenia, z kropka i malymi literami."""

    account_number: str
    """Testowy numer rachunku w zapisie ze spacjami."""

    account_documents: int
    """Liczba dokumentow zawierajacych numer rachunku w dowolnym wariancie zapisu."""

    broken_files: int
    """Pliki celowo niemozliwe do zaindeksowania: uszkodzony, pusty i zabezpieczony."""

    unsupported_files: int
    """Pliki o formacie nieobslugiwanym przez aplikacje."""

    scan_files: int
    """Pliki bez warstwy tekstowej, wymagajace OCR."""

    expected_queries: list[DemoQuery] = field(default_factory=list)
    """Zapytania referencyjne pokrywajace trzy tryby wyszukiwania."""

    def path_of(self, relative: str) -> Path:
        """Sciezka bezwzgledna dokumentu o podanej sciezce wzglednej."""
        return self.root.joinpath(*relative.split("/"))

    def queries_for_mode(self, mode: str) -> list[DemoQuery]:
        """Zapytania referencyjne w podanym trybie."""
        return [q for q in self.expected_queries if q.mode == mode]


# --- pomocnicze zapisy plikow ---------------------------------------------------


def _write_text(path: Path, lines: list[str], *, encoding: str = "utf-8") -> None:
    """Zapisuje plik tekstowy z zakonczeniami linii w stylu Windows."""
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\r\n".join(lines) + "\r\n", encoding=encoding)


def _write_csv(
    path: Path,
    header: list[str],
    rows: list[list[str]],
    *,
    encoding: str = "utf-8-sig",
    delimiter: str = ";",
) -> None:
    """Zapisuje plik CSV z separatorem srednika, jak eksport z polskiego Excela."""
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding=encoding, newline="") as handle:
        writer = csv.writer(handle, delimiter=delimiter, lineterminator="\r\n")
        writer.writerow(header)
        writer.writerows(rows)


def _write_docx(
    path: Path,
    title: str,
    paragraphs: list[str],
    *,
    author: str,
    created: _dt.datetime,
    table: tuple[list[str], list[list[str]]] | None = None,
) -> None:
    """Zapisuje dokument DOCX z naglowkiem, akapitami i opcjonalna tabela."""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - zaleznosc obowiazkowa projektu
        raise DependencyUnavailableError(
            "Biblioteka python-docx jest wymagana do wygenerowania zbioru demonstracyjnego.",
            cause=exc,
        ) from exc

    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    properties = document.core_properties
    properties.title = title
    properties.author = author
    properties.created = created
    properties.modified = created
    properties.language = "pl-PL"
    document.add_heading(title, level=1)
    for text in paragraphs:
        document.add_paragraph(text)
    if table is not None:
        header, rows = table
        grid = document.add_table(rows=1, cols=len(header))
        grid.style = "Table Grid"
        for column, label in enumerate(header):
            grid.rows[0].cells[column].text = label
        for row in rows:
            cells = grid.add_row().cells
            for column, value in enumerate(row):
                cells[column].text = value
    document.save(str(path))


def _write_xlsx(
    path: Path,
    sheets: list[tuple[str, list[str], list[list[str]]]],
    *,
    title: str,
    author: str,
) -> None:
    """Zapisuje skoroszyt XLSX z podanymi arkuszami."""
    try:
        from openpyxl import Workbook
    except ImportError as exc:  # pragma: no cover - zaleznosc obowiazkowa projektu
        raise DependencyUnavailableError(
            "Biblioteka openpyxl jest wymagana do wygenerowania zbioru demonstracyjnego.",
            cause=exc,
        ) from exc

    path.parent.mkdir(parents=True, exist_ok=True)
    book = Workbook()
    first: Any = book.active
    for index, (name, header, rows) in enumerate(sheets):
        sheet: Any
        if index == 0 and first is not None:
            sheet = first
            sheet.title = name
        else:
            sheet = book.create_sheet(title=name)
        sheet.append(header)
        for row in rows:
            sheet.append(row)
        for column in range(1, len(header) + 1):
            sheet.column_dimensions[sheet.cell(row=1, column=column).column_letter].width = 26
    book.properties.title = title
    book.properties.creator = author
    book.save(str(path))


# --- RTF ------------------------------------------------------------------------


def _rtf_escape(text: str) -> str:
    """Zamienia tekst na sekwencje RTF w stronie kodowej cp1250."""
    out: list[str] = []
    for char in text:
        if char in "\\{}":
            out.append("\\" + char)
        elif ord(char) < 128:
            out.append(char)
        else:
            try:
                encoded = char.encode("cp1250")
            except UnicodeEncodeError:
                out.append("?")
                continue
            out.append(f"\\'{encoded[0]:02x}")
    return "".join(out)


def _write_rtf(path: Path, title: str, paragraphs: list[str]) -> None:
    """Zapisuje dokument RTF z jawna deklaracja strony kodowej 1250."""
    path.parent.mkdir(parents=True, exist_ok=True)
    parts: list[str] = [
        r"{\rtf1\ansi\ansicpg1250\deff0\uc1",
        r"{\fonttbl{\f0\fswiss\fcharset238 Arial;}}",
        r"{\info{\title " + _rtf_escape(title) + r"}{\author Zespol FindDocs}}",
        r"\viewkind4\pard\f0\fs24 ",
        r"\b " + _rtf_escape(title) + r"\b0\par",
    ]
    parts.extend(_rtf_escape(paragraph) + r"\par" for paragraph in paragraphs)
    parts.append("}")
    path.write_bytes("\r\n".join(parts).encode("cp1250", errors="replace"))


# --- PDF ------------------------------------------------------------------------

#: Polskie znaki wraz z nazwami glifow, kodowane od pozycji 0x80.
_POLISH_GLYPHS: Final[tuple[tuple[str, str], ...]] = (
    ("ą", "aogonek"),
    ("ć", "cacute"),
    ("ę", "eogonek"),
    ("ł", "lslash"),
    ("ń", "nacute"),
    ("ó", "oacute"),
    ("ś", "sacute"),
    ("ź", "zacute"),
    ("ż", "zdotaccent"),
    ("Ą", "Aogonek"),
    ("Ć", "Cacute"),
    ("Ę", "Eogonek"),
    ("Ł", "Lslash"),
    ("Ń", "Nacute"),
    ("Ó", "Oacute"),
    ("Ś", "Sacute"),
    ("Ź", "Zacute"),
    ("Ż", "Zdotaccent"),
)

#: Pierwszy kod przypisany polskim znakom w tablicy Differences.
_FIRST_POLISH_CODE: Final = 0x80

#: Mapa znak polski -> kod bajtu w strumieniu tekstu PDF.
_PDF_CODE_BY_CHAR: Final[dict[str, int]] = {
    char: _FIRST_POLISH_CODE + index for index, (char, _name) in enumerate(_POLISH_GLYPHS)
}

#: Rozmiar strony A4 w punktach typograficznych.
_PDF_PAGE_WIDTH: Final = 595
_PDF_PAGE_HEIGHT: Final = 842

#: Ustawienia lamania tekstu na stronie PDF.
_PDF_LINE_HEIGHT: Final = 15
_PDF_TOP: Final = 790
_PDF_LEFT: Final = 56
_PDF_LINES_PER_PAGE: Final = 46
_PDF_WRAP_WIDTH: Final = 92

#: Znak z obszaru prywatnego zastepujacy spacje, ktorych nie wolno zlamac.
_KEEP_TOGETHER: Final = chr(0xE000)

#: Frazy, ktore musza zostac w jednej linii, zeby zapytania dokladne dzialaly.
_UNBREAKABLE: Final[tuple[str, ...]] = (
    ACCOUNT_SPACED,
    ACCOUNT_B_SPACED,
    ACCOUNT_TECH_SPACED,
)


def _pdf_encode_line(text: str) -> bytes:
    """Koduje jedna linie tekstu do postaci literalu PDF."""
    out = bytearray()
    for char in text:
        code = _PDF_CODE_BY_CHAR.get(char)
        if code is None:
            point = ord(char)
            code = point if 32 <= point < 127 else ord("?")
        if code in (0x28, 0x29, 0x5C):
            out.append(0x5C)
        out.append(code)
    return bytes(out)


def _pdf_text_string(text: str) -> bytes:
    """Koduje napis metadanych jako szesnastkowy ciag UTF-16BE."""
    raw = b"\xfe\xff" + text.encode("utf-16-be")
    return b"<" + raw.hex().upper().encode("ascii") + b">"


def _pdf_date(moment: _dt.datetime) -> str:
    """Data w formacie akceptowanym przez czytniki PDF."""
    return moment.strftime("D:%Y%m%d%H%M%S+02'00'")


def _pdf_stream_object(data: bytes, extra: str = "") -> bytes:
    """Buduje obiekt strumienia o podanej zawartosci."""
    header = f"<< /Length {len(data)}{extra} >>\nstream\n".encode("latin-1")
    return header + data + b"\nendstream"


def _pdf_tounicode_cmap() -> bytes:
    """Mapa ToUnicode pozwalajaca odczytac tekst wraz z polskimi znakami."""
    entries: list[tuple[int, str]] = [(code, chr(code)) for code in range(32, 127)]
    entries.extend((code, char) for char, code in _PDF_CODE_BY_CHAR.items())
    lines: list[str] = [
        "/CIDInit /ProcSet findresource begin",
        "12 dict begin",
        "begincmap",
        "/CIDSystemInfo << /Registry (Adobe) /Ordering (UCS) /Supplement 0 >> def",
        "/CMapName /Adobe-Identity-UCS def",
        "/CMapType 2 def",
        "1 begincodespacerange",
        "<00> <FF>",
        "endcodespacerange",
    ]
    for start in range(0, len(entries), 100):
        block = entries[start : start + 100]
        lines.append(f"{len(block)} beginbfchar")
        lines.extend(f"<{code:02X}> <{ord(char):04X}>" for code, char in block)
        lines.append("endbfchar")
    lines.extend(
        [
            "endcmap",
            "CMapName currentdict /CMap defineresource pop",
            "end",
            "end",
        ]
    )
    return "\n".join(lines).encode("latin-1")


def _pdf_encoding_object() -> bytes:
    """Slownik kodowania z tablica Differences dla polskich glifow."""
    names = " ".join(f"/{name}" for _char, name in _POLISH_GLYPHS)
    return (
        "<< /Type /Encoding /BaseEncoding /WinAnsiEncoding "
        f"/Differences [ {_FIRST_POLISH_CODE} {names} ] >>"
    ).encode("latin-1")


def _assemble_pdf(objects: list[bytes], *, root: int, info: int | None = None) -> bytes:
    """Sklada dokument PDF z listy cial obiektow ponumerowanych od jedynki."""
    out = bytearray(b"%PDF-1.7\n%\xe2\xe3\xcf\xd3\n")
    offsets: list[int] = []
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{number} 0 obj\n".encode("latin-1")
        out += body
        out += b"\nendobj\n"
    xref_offset = len(out)
    size = len(objects) + 1
    out += f"xref\n0 {size}\n".encode("latin-1")
    out += b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode("latin-1")
    trailer = f"<< /Size {size} /Root {root} 0 R"
    if info is not None:
        trailer += f" /Info {info} 0 R"
    trailer += " >>"
    out += b"trailer\n" + trailer.encode("latin-1") + b"\n"
    out += f"startxref\n{xref_offset}\n".encode("latin-1")
    out += b"%%EOF\n"
    return bytes(out)


def _pdf_wrap(paragraphs: list[str]) -> list[str]:
    """Lamie akapity na linie mieszczace sie na stronie.

    Numery rachunkow sa chronione przed podzialem, zeby zapytania dokladne
    trafialy na ciagly zapis numeru w jednej linii.
    """
    lines: list[str] = []
    for paragraph in paragraphs:
        if not paragraph:
            lines.append("")
            continue
        guarded = paragraph
        for phrase in _UNBREAKABLE:
            guarded = guarded.replace(phrase, phrase.replace(" ", _KEEP_TOGETHER))
        wrapped = textwrap.wrap(guarded, width=_PDF_WRAP_WIDTH)
        lines.extend(line.replace(_KEEP_TOGETHER, " ") for line in wrapped or [""])
        lines.append("")
    return lines


def _pdf_page_content(lines: list[str]) -> bytes:
    """Strumien tresci jednej strony tekstowej."""
    out = bytearray(b"BT\n/F1 11 Tf\n")
    out += f"{_PDF_LINE_HEIGHT} TL\n{_PDF_LEFT} {_PDF_TOP} Td\n".encode("latin-1")
    for line in lines:
        out += b"(" + _pdf_encode_line(line) + b") Tj\nT*\n"
    out += b"ET\n"
    return bytes(out)


def build_text_pdf(title: str, paragraphs: list[str], *, created: _dt.datetime) -> bytes:
    """Buduje jednoczesciowy dokument PDF z warstwa tekstowa."""
    lines = _pdf_wrap(paragraphs)
    pages = [
        lines[start : start + _PDF_LINES_PER_PAGE]
        for start in range(0, max(len(lines), 1), _PDF_LINES_PER_PAGE)
    ]
    objects: list[bytes] = [
        b"",
        b"",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica "
        b"/FirstChar 32 /LastChar 255 /Encoding 5 0 R /ToUnicode 4 0 R >>",
        _pdf_stream_object(_pdf_tounicode_cmap()),
        _pdf_encoding_object(),
    ]
    kids: list[int] = []
    for page_lines in pages:
        objects.append(_pdf_stream_object(_pdf_page_content(page_lines)))
        content_number = len(objects)
        page = (
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 {_PDF_PAGE_WIDTH} {_PDF_PAGE_HEIGHT}] "
            f"/Resources << /Font << /F1 3 0 R >> >> /Contents {content_number} 0 R >>"
        )
        objects.append(page.encode("latin-1"))
        kids.append(len(objects))
    objects[0] = b"<< /Type /Catalog /Pages 2 0 R >>"
    kid_refs = " ".join(f"{number} 0 R" for number in kids)
    objects[1] = f"<< /Type /Pages /Kids [ {kid_refs} ] /Count {len(kids)} >>".encode("latin-1")
    info_body = bytearray(b"<< /Title ")
    info_body += _pdf_text_string(title)
    info_body += b" /Author "
    info_body += _pdf_text_string("Zespol FindDocs")
    info_body += b" /Producer "
    info_body += _pdf_text_string("FindDocs demo")
    info_body += f" /CreationDate ({_pdf_date(created)}) >>".encode("latin-1")
    objects.append(bytes(info_body))
    return _assemble_pdf(objects, root=1, info=len(objects))


def build_image_pdf(jpeg: bytes, *, width: int, height: int) -> bytes:
    """Buduje dokument PDF z osadzonym obrazem JPEG, bez warstwy tekstowej."""
    content = (f"q {_PDF_PAGE_WIDTH} 0 0 {_PDF_PAGE_HEIGHT} 0 0 cm /Im1 Do Q\n").encode("latin-1")
    image_dict = (
        f"<< /Type /XObject /Subtype /Image /Width {width} /Height {height} "
        f"/ColorSpace /DeviceRGB /BitsPerComponent 8 /Filter /DCTDecode /Length {len(jpeg)} >>"
    )
    objects: list[bytes] = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [ 3 0 R ] /Count 1 >>",
        (
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 {_PDF_PAGE_WIDTH} {_PDF_PAGE_HEIGHT}] "
            "/Resources << /XObject << /Im1 5 0 R >> >> /Contents 4 0 R >>"
        ).encode("latin-1"),
        _pdf_stream_object(content),
        image_dict.encode("latin-1") + b"\nstream\n" + jpeg + b"\nendstream",
    ]
    return _assemble_pdf(objects, root=1)


# --- obrazy skanow --------------------------------------------------------------

#: Czcionki systemowe probowane przy renderowaniu skanow.
_FONT_CANDIDATES: Final[tuple[str, ...]] = (
    "C:/Windows/Fonts/arial.ttf",
    "C:/Windows/Fonts/calibri.ttf",
    "C:/Windows/Fonts/segoeui.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
)


def _load_font(size: int) -> Any:
    """Zwraca czcionke o zadanym rozmiarze, z awaryjnym powrotem do czcionki wbudowanej."""
    from PIL import ImageFont

    for candidate in _FONT_CANDIDATES:
        path = Path(candidate)
        if not path.exists():
            continue
        try:
            return ImageFont.truetype(str(path), size)
        except OSError:  # pragma: no cover - uszkodzony plik czcionki
            continue
    try:
        return ImageFont.load_default(size=size)
    except TypeError:  # pragma: no cover - starsze wersje Pillow
        return ImageFont.load_default()


def render_scan_image(
    lines: list[str],
    *,
    width: int = 1240,
    height: int = 1754,
    rotation: float = 0.4,
) -> Any:
    """Renderuje obraz udajacy skan dokumentu: czarny tekst na jasnym tle."""
    try:
        from PIL import Image, ImageDraw
    except ImportError as exc:  # pragma: no cover - zaleznosc obowiazkowa projektu
        raise DependencyUnavailableError(
            "Biblioteka Pillow jest wymagana do wygenerowania skanow demonstracyjnych.",
            cause=exc,
        ) from exc

    image = Image.new("RGB", (width, height), (248, 246, 240))
    draw = ImageDraw.Draw(image)
    draw.rectangle([40, 40, width - 40, height - 40], outline=(190, 186, 178), width=3)
    title_font = _load_font(52)
    body_font = _load_font(36)
    position = 120
    for index, line in enumerate(lines):
        font = title_font if index == 0 else body_font
        draw.text((100, position), line, fill=(24, 24, 28), font=font)
        position += 78 if index == 0 else 56
    rotated = image.rotate(rotation, expand=False, fillcolor=(248, 246, 240))
    return rotated


def _write_scan_png(path: Path, lines: list[str]) -> None:
    """Zapisuje skan jako plik PNG."""
    path.parent.mkdir(parents=True, exist_ok=True)
    image = render_scan_image(lines)
    image.save(str(path), format="PNG", dpi=(150, 150))


def _write_scan_pdf(path: Path, lines: list[str]) -> None:
    """Zapisuje skan jako PDF z obrazem JPEG i bez warstwy tekstowej."""
    path.parent.mkdir(parents=True, exist_ok=True)
    image = render_scan_image(lines, width=1000, height=1414, rotation=-0.5)
    buffer = io.BytesIO()
    image.save(buffer, format="JPEG", quality=85, optimize=True)
    jpeg = buffer.getvalue()
    path.write_bytes(build_image_pdf(jpeg, width=image.width, height=image.height))


# --- wiadomosci e-mail ----------------------------------------------------------


def _write_eml(
    path: Path,
    *,
    sender: str,
    recipients: list[str],
    subject: str,
    moment: _dt.datetime,
    body: list[str],
    message_id: str,
    attachment: tuple[str, bytes] | None = None,
) -> None:
    """Zapisuje wiadomosc EML, opcjonalnie z zalacznikiem PDF."""
    path.parent.mkdir(parents=True, exist_ok=True)
    message = EmailMessage()
    message["From"] = sender
    message["To"] = ", ".join(recipients)
    message["Subject"] = subject
    message["Date"] = format_datetime(moment)
    message["Message-ID"] = message_id
    message.set_content("\r\n".join(body), subtype="plain", charset="utf-8", cte="quoted-printable")
    if attachment is not None:
        name, data = attachment
        message.add_attachment(
            data,
            maintype="application",
            subtype="pdf",
            filename=name,
        )
        # Granica czesci jest losowa, wiec ustawiamy wlasna, zeby zbior byl powtarzalny.
        message.set_boundary(f"=====granica-{moment.strftime('%Y%m%d%H%M%S')}=====")
    path.write_bytes(message.as_bytes())


# --- zaszyfrowane archiwum ZIP --------------------------------------------------


def _crc_table() -> tuple[int, ...]:
    """Tablica CRC-32 uzywana przez klasyczne szyfrowanie ZIP."""
    table: list[int] = []
    for index in range(256):
        value = index
        for _ in range(8):
            value = (value >> 1) ^ (0xEDB88320 if value & 1 else 0)
        table.append(value)
    return tuple(table)


_CRC_TABLE: Final[tuple[int, ...]] = _crc_table()


class _ZipCrypto:
    """Klasyczny strumieniowy szyfrator ZIP (PKWARE) uzyty do pliku demonstracyjnego."""

    __slots__ = ("_keys",)

    def __init__(self, phrase: bytes) -> None:
        self._keys: list[int] = [0x12345678, 0x23456789, 0x34567890]
        for byte in phrase:
            self._update(byte)

    def _update(self, byte: int) -> None:
        keys = self._keys
        keys[0] = (keys[0] >> 8) ^ _CRC_TABLE[(keys[0] ^ byte) & 0xFF]
        keys[1] = (keys[1] + (keys[0] & 0xFF)) & 0xFFFFFFFF
        keys[1] = (keys[1] * 134775813 + 1) & 0xFFFFFFFF
        top = (keys[1] >> 24) & 0xFF
        keys[2] = (keys[2] >> 8) ^ _CRC_TABLE[(keys[2] ^ top) & 0xFF]

    def _stream_byte(self) -> int:
        temp = (self._keys[2] | 2) & 0xFFFF
        return ((temp * (temp ^ 1)) >> 8) & 0xFF

    def encrypt(self, data: bytes) -> bytes:
        """Szyfruje bajty w kolejnosci wystepowania."""
        out = bytearray()
        for byte in data:
            out.append(byte ^ self._stream_byte())
            self._update(byte)
        return bytes(out)


def _dos_datetime(moment: _dt.datetime) -> tuple[int, int]:
    """Czas i data w formacie MS-DOS uzywanym w naglowkach ZIP."""
    dos_time = (moment.hour << 11) | (moment.minute << 5) | (moment.second // 2)
    dos_date = ((moment.year - 1980) << 9) | (moment.month << 5) | moment.day
    return dos_time, dos_date


def write_protected_zip(
    path: Path,
    entries: list[tuple[str, bytes]],
    *,
    phrase: str,
    rng: random.Random,
    moment: _dt.datetime,
) -> None:
    """Zapisuje archiwum ZIP z klasycznym szyfrowaniem kazdego wpisu."""
    path.parent.mkdir(parents=True, exist_ok=True)
    dos_time, dos_date = _dos_datetime(moment)
    payload = bytearray()
    central = bytearray()
    for name, data in entries:
        crc = zlib.crc32(data) & 0xFFFFFFFF
        compressor = zlib.compressobj(9, zlib.DEFLATED, -15)
        compressed = compressor.compress(data) + compressor.flush()
        cipher = _ZipCrypto(phrase.encode("cp1250"))
        preamble = bytes(rng.randrange(256) for _ in range(11)) + bytes([(crc >> 24) & 0xFF])
        encrypted = cipher.encrypt(preamble + compressed)
        offset = len(payload)
        raw_name = name.encode("ascii")
        payload += struct.pack(
            "<IHHHHHIIIHH",
            0x04034B50,
            20,
            0x0001,
            8,
            dos_time,
            dos_date,
            crc,
            len(encrypted),
            len(data),
            len(raw_name),
            0,
        )
        payload += raw_name + encrypted
        central += struct.pack(
            "<IHHHHHHIIIHHHHHII",
            0x02014B50,
            0x0314,
            20,
            0x0001,
            8,
            dos_time,
            dos_date,
            crc,
            len(encrypted),
            len(data),
            len(raw_name),
            0,
            0,
            0,
            0,
            0o644 << 16,
            offset,
        )
        central += raw_name
    directory_offset = len(payload)
    payload += central
    payload += struct.pack(
        "<IHHHHIIH",
        0x06054B50,
        0,
        0,
        len(entries),
        len(entries),
        len(central),
        directory_offset,
        0,
    )
    path.write_bytes(bytes(payload))


# --- tresc dokumentow -----------------------------------------------------------


def _build_procedury(root: Path) -> None:
    """Tworzy katalog z procedurami bankowymi."""
    _write_docx(
        root / "procedury" / "procedura-przelewow-2015.docx",
        "Procedura realizacji przelewów krajowych",
        [
            "Wersja 3.0. Procedura obowiązuje od 24.07.2015 r. (2015-07-24) i zastępuje "
            "wersję 2.4 z 12.01.2012 r.",
            "Dokument zatwierdził dyrektor operacyjny Kowalski Jan w dniu 24 lipca 2015 r.",
            "1. Polecenie przelewu przyjęte do godziny 15:30 jest realizowane w tym samym dniu "
            "roboczym, w ostatniej sesji rozliczeniowej.",
            "2. Dyspozycje złożone po godzinie 15:30 trafiają do pierwszej sesji rozliczeniowej "
            "następnego dnia roboczego.",
            "3. Przelew powyżej 50 000,00 zł wymaga autoryzacji dwóch pracowników zespołu "
            "rozliczeń.",
            "4. Przykładowy rachunek uznaniowy używany w testach: " + ACCOUNT_SPACED + ".",
            "5. Rachunek techniczny do rozliczeń wewnętrznych: " + ACCOUNT_TECH_SPACED + ".",
            "6. Reklamacje dotyczące przelewów rozpatruje zespół obsługi klienta zgodnie "
            "z odrębną procedurą.",
            "Właściciel procedury: Departament Rozliczeń, " + BANK + ".",
            DISCLAIMER,
        ],
        author=PERSON_A,
        created=_dt.datetime(2015, 7, 24, 9, 15),
        table=(
            ["Etap", "Termin", "Odpowiedzialny"],
            [
                ["Przyjęcie dyspozycji", "do 15:30", "Zespół obsługi klienta"],
                ["Weryfikacja", "do 16:00", PERSON_C],
                ["Sesja rozliczeniowa", "16:30", "Departament Rozliczeń"],
            ],
        ),
    )
    _write_docx(
        root / "procedury" / "procedura-przeciwdzialania-praniu-pieniedzy.docx",
        "Procedura przeciwdziałania praniu pieniędzy",
        [
            "Celem procedury jest opisanie zasad poznania klienta oraz badania powiązań "
            "pomiędzy podmiotami obsługiwanymi przez bank.",
            "Analityk ustala strukturę właścicielską, beneficjenta rzeczywistego oraz źródło "
            "pochodzenia majątku przed nawiązaniem współpracy.",
            "Powiązania kapitałowe i osobowe pomiędzy klientami są odnotowywane w rejestrze "
            "powiązań i przeglądane co kwartał.",
            "Przykład zarejestrowanego powiązania: " + CLIENT_A + " występuje jako poręczyciel "
            "zobowiązań spółki " + CLIENT_B + ".",
            "Rachunek objęty wzmożonym monitoringiem: " + ACCOUNT_DASHED + ".",
            "Transakcje o nietypowej częstotliwości są raportowane w terminie dwóch dni "
            "roboczych od ich wykrycia.",
            "Osoba odpowiedzialna za rejestr powiązań: " + PERSON_B + ".",
            DISCLAIMER,
        ],
        author=PERSON_B,
        created=_dt.datetime(2016, 3, 2, 11, 40),
    )
    _write_docx(
        root / "procedury" / "procedura-obslugi-reklamacji.docx",
        "Procedura obsługi reklamacji klientów",
        [
            "Reklamacja może dotyczyć nieuznanej wpłaty, podwójnego obciążenia rachunku albo "
            "kwestionowanej transakcji kartowej.",
            "Termin rozpatrzenia reklamacji wynosi 30 dni kalendarzowych od dnia jej "
            "przyjęcia w oddziale albo drogą elektroniczną.",
            "Przykład zgłoszenia: klient kwestionuje płatność kartą " + CARD_MASKED + " na "
            "kwotę 314 zł wykonaną w hurtowni biurowej.",
            "W systemie transakcja widnieje jako obciążenie 314,00 PLN i wymaga porównania "
            "z zapisem terminala.",
            "Rachunek klienta zgłaszającego reklamację: " + ACCOUNT_COMPACT + ".",
            "Zgłoszenie prowadzi " + PERSON_C + " z zespołu obsługi klienta.",
            "Odmowa uznania reklamacji wymaga pisemnego uzasadnienia i pouczenia o trybie "
            "odwoławczym.",
            DISCLAIMER,
        ],
        author=PERSON_C,
        created=_dt.datetime(2017, 9, 11, 8, 5),
    )
    _write_text(
        root / "procedury" / "instrukcja-ksiegowania-2007.txt",
        [
            "INSTRUKCJA KSIĘGOWANIA OPERACJI PO AWARII SYSTEMU",
            "=================================================",
            "",
            "Data zdarzenia: 05.05.2007 (2007-05-05), sobota.",
            "W dniu 5 maja 2007 r. o godzinie 04:12 doszło do awarii systemu księgowego.",
            "Sesja rozliczeniowa z tego dnia została wstrzymana na cztery godziny.",
            "",
            "Zakres prac:",
            "1. Operacje z 05.05.2007 zostały zaksięgowane ręcznie przez zespół rozliczeń.",
            "2. Uzgodnienie sald wykonano 07.05.2007 na rachunku " + ACCOUNT_COMPACT + ".",
            "3. Różnice kursowe rozliczono na rachunku technicznym " + ACCOUNT_TECH_SPACED + ".",
            "4. Protokół z awarii podpisał " + PERSON_D + ".",
            "",
            "Uwagi: dokumentacja papierowa z 5 maja 2007 r. została zarchiwizowana",
            "w segregatorze o sygnaturze 00/ARCH/2007/44.",
            "",
            DISCLAIMER,
        ],
    )
    _write_text(
        root / "procedury" / "procedura-archiwizacji.txt",
        [
            "PROCEDURA ARCHIWIZACJI DOKUMENTACJI PAPIEROWEJ I ELEKTRONICZNEJ",
            "",
            "1. Dokumenty rozliczeniowe przechowuje się przez pięć lat od zakończenia roku",
            "   obrotowego, w którym powstały.",
            "2. Umowy rachunków bankowych przechowuje się przez dziesięć lat od rozwiązania",
            "   umowy.",
            "3. Nośniki elektroniczne są kopiowane raz w miesiącu, a kopia jest przechowywana",
            "   w innej lokalizacji niż oryginał.",
            "4. Zniszczenie dokumentacji wymaga protokołu podpisanego przez dwie osoby.",
            "5. Akta rachunku " + ACCOUNT_SPACED + " są przechowywane w archiwum zakładowym",
            "   pod sygnaturą 00/ARCH/2015/12.",
            "6. Wnioski o udostępnienie akt rozpatruje " + PERSON_D + " w terminie pięciu dni.",
            "",
            DISCLAIMER,
        ],
    )


def _build_umowy(root: Path) -> None:
    """Tworzy katalog z umowami."""
    _write_docx(
        root / "umowy" / "umowa-rachunku-acme.docx",
        f"Umowa rachunku bieżącego nr {CONTRACT_ACCOUNT}",
        [
            f"Umowa zawarta w dniu 12.04.2005 r. pomiędzy {BANK} a spółką {CLIENT_A}.",
            "Spółkę reprezentuje prezes zarządu " + PERSON_A + ", legitymujący się "
            "pełnomocnictwem nr 00/PEL/2005/9.",
            "Bank otwiera i prowadzi rachunek bieżący o numerze " + ACCOUNT_SPACED + ".",
            "Opłata za prowadzenie rachunku wynosi 29,00 zł miesięcznie i jest pobierana "
            "ostatniego dnia roboczego miesiąca.",
            "Zlecenia płatnicze są realizowane zgodnie z procedurą realizacji przelewów "
            "krajowych obowiązującą w banku.",
            "Umowa została zmieniona aneksem nr 1 z dnia 05.05.2007 r.",
            DISCLAIMER,
        ],
        author=PERSON_A,
        created=_dt.datetime(2005, 4, 12, 10, 0),
        table=(
            ["Pozycja", "Wartość"],
            [
                ["Numer umowy", CONTRACT_ACCOUNT],
                ["Numer rachunku", ACCOUNT_SPACED],
                ["Posiadacz", CLIENT_A],
                ["Data zawarcia", "12.04.2005"],
            ],
        ),
    )
    _write_docx(
        root / "umowy" / "umowa-kredytowa-nowak-bud.docx",
        f"Umowa kredytu obrotowego nr {CONTRACT_CREDIT}",
        [
            f"Kredytobiorca: {CLIENT_B}, wpisana do rejestru przedsiębiorców pod numerem "
            "00-B-9902.",
            "Kwota kredytu: 250 000,00 zł, okres kredytowania: 36 miesięcy.",
            f"Poręczycielem zobowiązania jest {CLIENT_A}, co stanowi powiązanie gospodarcze "
            "pomiędzy oboma klientami banku.",
            "Spłata rat następuje z rachunku poręczyciela " + ACCOUNT_DASHED + " w przypadku "
            "braku środków na rachunku kredytobiorcy " + ACCOUNT_B_SPACED + ".",
            "Umowę w imieniu kredytobiorcy podpisał "
            + PERSON_D
            + ", a ze strony banku "
            + PERSON_C
            + ".",
            "Zabezpieczeniem kredytu jest weksel własny in blanco wraz z deklaracją wekslową.",
            DISCLAIMER,
        ],
        author=PERSON_C,
        created=_dt.datetime(2015, 6, 18, 13, 30),
    )
    _write_rtf(
        root / "umowy" / "aneks-nr-1-do-umowy-rachunku.rtf",
        f"Aneks nr 1 do umowy rachunku bieżącego nr {CONTRACT_ACCOUNT}",
        [
            "Aneks zawarto w dniu 05.05.2007 r. (5 maja 2007 r.) w siedzibie banku.",
            f"Strony: {BANK} oraz {CLIENT_A}, reprezentowana przez {PERSON_A}.",
            "Strony zgodnie zmieniają paragraf 4 umowy w ten sposób, że opłata za prowadzenie "
            "rachunku " + ACCOUNT_DASHED + " zostaje obniżona do 19,00 zł miesięcznie.",
            "Zmiana obowiązuje od pierwszego dnia miesiąca następującego po dniu podpisania "
            "aneksu.",
            "Pozostałe postanowienia umowy pozostają bez zmian.",
            "Aneks sporządzono w dwóch jednobrzmiących egzemplarzach.",
            DISCLAIMER,
        ],
    )
    pdf = build_text_pdf(
        "Porozumienie rozliczeniowe",
        [
            "POROZUMIENIE ROZLICZENIOWE",
            "",
            f"zawarte w dniu 24 lipca 2015 r. (24.07.2015) pomiędzy {CLIENT_A} a {CLIENT_B}",
            "",
            f"1. Strony ustalają, że rozliczenia wzajemne są prowadzone przez rachunek "
            f"{ACCOUNT_SPACED} prowadzony przez {BANK}",
            f"2. Podstawą rozliczenia jest faktura {INVOICE_ID} wystawiona za usługi "
            "wykonane w lipcu 2015 r.",
            "3. Kwota bezsporna wynosi 12 400,00 zł i zostanie zapłacona w terminie siedmiu dni.",
            "4. Kwota sporna 314 zł dotyczy opłaty dodatkowej i podlega dalszym uzgodnieniom.",
            "5. Porozumienie nie zmienia zabezpieczeń ustanowionych w umowie kredytu "
            f"{CONTRACT_CREDIT}.",
            "",
            f"Podpisali: {PERSON_A} oraz {PERSON_D}.",
            "",
            DISCLAIMER,
        ],
        created=_dt.datetime(2015, 7, 24, 12, 0),
    )
    target = root / "umowy" / "porozumienie-rozliczeniowe.pdf"
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_bytes(pdf)


def _disclaimer_row(columns: int) -> list[str]:
    """Wiersz stopki tabeli z informacja o fikcyjnosci danych."""
    row = [""] * columns
    row[1 if columns > 1 else 0] = DISCLAIMER
    return row


def _transaction_columns() -> list[str]:
    """Naglowki plikow z transakcjami."""
    return ["data", "opis", "kwota", "waluta", "rachunek", "kontrahent"]


def _build_transakcje(root: Path, rng: random.Random) -> None:
    """Tworzy katalogi z transakcjami obu klientow."""
    header = _transaction_columns()
    rows_a_2015: list[list[str]] = [
        [
            "2015-07-20",
            f"Przelew wychodzący, faktura {INVOICE_ID}",
            "-12 400,00",
            "PLN",
            ACCOUNT_SPACED,
            CLIENT_B,
        ],
        [
            "2015-07-22",
            f"Płatność kartą {CARD_MASKED}, hurtownia biurowa",
            "-314,00",
            "PLN",
            ACCOUNT_SPACED,
            "Biuro Serwis sp. z o.o.",
        ],
        [
            "2015-07-24",
            "Polecenie przelewu, sesja poranna",
            "-8 900,50",
            "PLN",
            ACCOUNT_SPACED,
            CLIENT_B,
        ],
        [
            "2015-07-24",
            "Wpływ z tytułu rozliczenia porozumienia",
            "24 750,00",
            "PLN",
            ACCOUNT_SPACED,
            "Zakład Usług Miejskich sp. z o.o.",
        ],
        [
            "2015-07-31",
            "Opłata za prowadzenie rachunku",
            "-29,00",
            "PLN",
            ACCOUNT_SPACED,
            BANK,
        ],
    ]
    for day in (21, 23, 27, 28, 29):
        amount = rng.randrange(120, 9800) + rng.randrange(0, 100) / 100
        rows_a_2015.append(
            [
                f"2015-07-{day:02d}",
                rng.choice(
                    [
                        "Przelew wychodzący, rozliczenie usług",
                        "Wpłata gotówkowa w oddziale",
                        "Przelew przychodzący od kontrahenta",
                        "Opłata za przelew natychmiastowy",
                    ]
                ),
                f"{amount:,.2f}".replace(",", " ").replace(".", ","),
                "PLN",
                ACCOUNT_SPACED,
                rng.choice(
                    [
                        "Drukarnia Fikcyjna sp. z o.o.",
                        "Transport Wisła sp. z o.o.",
                        CLIENT_B,
                        "Serwis Biurowy Zieliński",
                    ]
                ),
            ]
        )
    rows_a_2015.append(_disclaimer_row(len(header)))
    _write_csv(root / "transakcje" / "klientA" / "transakcje-2015-07.csv", header, rows_a_2015)

    rows_a_2007: list[list[str]] = [
        [
            "04.05.2007",
            "Przelew przychodzący, zaliczka na roboty",
            "3 400,00",
            "PLN",
            ACCOUNT_COMPACT,
            CLIENT_B,
        ],
        [
            "05.05.2007",
            "Księgowanie ręczne po awarii systemu",
            "1 250,00",
            "PLN",
            ACCOUNT_COMPACT,
            BANK,
        ],
        [
            "05.05.2007",
            "Wpłata gotówkowa w oddziale, potwierdzenie 00/KP/2007/318",
            "314,00",
            "PLN",
            ACCOUNT_COMPACT,
            CLIENT_A,
        ],
        [
            "07.05.2007",
            "Uzgodnienie sald po awarii, korekta techniczna",
            "-1 250,00",
            "PLN",
            ACCOUNT_COMPACT,
            BANK,
        ],
        [
            "09.05.2007",
            "Opłata za prowadzenie rachunku",
            "-19,00",
            "PLN",
            ACCOUNT_COMPACT,
            BANK,
        ],
    ]
    rows_a_2007.append(_disclaimer_row(len(header)))
    _write_csv(
        root / "transakcje" / "klientA" / "transakcje-2007-05.csv",
        header,
        rows_a_2007,
        encoding="cp1250",
    )

    rows_b_2015: list[list[str]] = [
        [
            "2015-07-20",
            f"Wpływ z rachunku {ACCOUNT_SPACED}, faktura {INVOICE_ID}",
            "12 400,00",
            "PLN",
            ACCOUNT_B_SPACED,
            CLIENT_A,
        ],
        [
            "2015-07-24",
            f"Wpływ z rachunku {ACCOUNT_SPACED}, zaliczka",
            "8 900,50",
            "PLN",
            ACCOUNT_B_SPACED,
            CLIENT_A,
        ],
        [
            "2015-07-25",
            "Wypłata wynagrodzeń, lista płac 99/LP/2015/07",
            "-31 200,00",
            "PLN",
            ACCOUNT_B_SPACED,
            "Pracownicy",
        ],
        [
            "2015-07-28",
            "Rata kredytu obrotowego " + CONTRACT_CREDIT,
            "-7 100,00",
            "PLN",
            ACCOUNT_B_SPACED,
            BANK,
        ],
    ]
    for day in (21, 26, 30):
        amount = rng.randrange(300, 15000) + rng.randrange(0, 100) / 100
        rows_b_2015.append(
            [
                f"2015-07-{day:02d}",
                rng.choice(
                    [
                        "Zakup materiałów budowlanych",
                        "Przelew do podwykonawcy",
                        "Wpłata gotówkowa w oddziale",
                    ]
                ),
                f"-{amount:,.2f}".replace(",", " ").replace(".", ","),
                "PLN",
                ACCOUNT_B_SPACED,
                rng.choice(
                    [
                        "Skład Budowlany Fikcja sp. z o.o.",
                        "Usługi Ziemne Dąbrowski",
                        "Hurtownia Stali Nowa sp. z o.o.",
                    ]
                ),
            ]
        )
    rows_b_2015.append(_disclaimer_row(len(header)))
    _write_csv(root / "transakcje" / "klientB" / "transakcje-2015-07.csv", header, rows_b_2015)

    kontrahenci = [
        [
            CLIENT_B,
            "00-B-9902",
            ACCOUNT_B_SPACED,
            "podmiot zestawienia",
            "2005-04-12",
        ],
        [
            CLIENT_A,
            "00-A-4471",
            ACCOUNT_COMPACT,
            "poręczyciel kredytu " + CONTRACT_CREDIT,
            "2015-06-18",
        ],
        [
            "Skład Budowlany Fikcja sp. z o.o.",
            "99-K-1002",
            "99 1000 2000 3000 4000 5000 6000",
            "dostawca materiałów",
            "2014-02-03",
        ],
        [
            "Usługi Ziemne Dąbrowski",
            "99-K-1188",
            "99 2000 3000 4000 5000 6000 7000",
            "podwykonawca",
            "2015-03-17",
        ],
    ]
    _write_xlsx(
        root / "transakcje" / "klientB" / "zestawienie-klientB.xlsx",
        [
            (
                "Transakcje",
                _transaction_columns(),
                [row[:] for row in rows_b_2015],
            ),
            (
                "Kontrahenci",
                ["nazwa", "identyfikator", "rachunek", "rodzaj powiązania", "data ustalenia"],
                [*kontrahenci, [DISCLAIMER, "", "", "", ""]],
            ),
        ],
        title=f"Zestawienie klienta B, {CLIENT_B}",
        author=PERSON_D,
    )


def _build_korespondencja(root: Path) -> None:
    """Tworzy katalog z korespondencja, w tym wiadomosc z zalacznikiem PDF."""
    attachment = build_text_pdf(
        "Potwierdzenie wykonania przelewu",
        [
            "POTWIERDZENIE WYKONANIA PRZELEWU",
            "",
            "Data realizacji: 24.07.2015 (2015-07-24), sesja poranna.",
            f"Rachunek obciążany: {ACCOUNT_SPACED}.",
            f"Odbiorca: {CLIENT_B}, rachunek {ACCOUNT_B_SPACED}.",
            "Kwota: 8 900,50 zł.",
            f"Tytuł: zaliczka do faktury {INVOICE_ID}.",
            f"Dokument wystawił automatycznie system {BANK}.",
            "",
            DISCLAIMER,
        ],
        created=_dt.datetime(2015, 7, 24, 10, 5),
    )
    _write_eml(
        root / "korespondencja" / "potwierdzenie-przelewu.eml",
        sender=f"{BANK} <powiadomienia@bank-testowy.example>",
        recipients=[f"Księgowość {CLIENT_A} <ksiegowosc@acme-polska.example>"],
        subject="Potwierdzenie przelewu z dnia 24.07.2015",
        moment=_dt.datetime(2015, 7, 24, 10, 12, tzinfo=_dt.timezone(_dt.timedelta(hours=2))),
        body=[
            "Dzień dobry,",
            "",
            "w załączeniu przesyłamy potwierdzenie przelewu",
            "wykonanego 24 lipca 2015 r. (2015-07-24).",
            "",
            f"Rachunek obciążany: {ACCOUNT_SPACED}",
            f"Odbiorca: {CLIENT_B}",
            "Kwota: 8 900,50 zł",
            f"Tytuł: zaliczka do faktury {INVOICE_ID}",
            "",
            "Pozdrawiamy",
            f"Zespół Obsługi Klienta, {BANK}",
            "",
            DISCLAIMER,
        ],
        message_id="<demo-0001@finddocs.example>",
        attachment=("potwierdzenie-przelewu-2015-07-24.pdf", attachment),
    )
    _write_eml(
        root / "korespondencja" / "zapytanie-o-ksiegowania.eml",
        sender=f"{PERSON_A} <jan.kowalski@acme-polska.example>",
        recipients=[f"{PERSON_B} <anna.nowak@bank-testowy.example>"],
        subject="Zapytanie o księgowania z 05.05.2007",
        moment=_dt.datetime(2007, 5, 6, 8, 41, tzinfo=_dt.timezone(_dt.timedelta(hours=2))),
        body=[
            "Dzień dobry Pani Anno,",
            "",
            "proszę o wyjaśnienie księgowań z 05.05.2007 (2007-05-05).",
            "W dniu 5 maja 2007 r. na wyciągu widnieją dwie pozycje,",
            "których nie potrafimy uzgodnić z naszą ewidencją.",
            "",
            f"Rachunek: {ACCOUNT_DASHED}",
            "Kwoty: 1 250,00 zł oraz 314,00 zł",
            "",
            f"Dodatkowo proszę o informację, czy powiązanie spółki {CLIENT_A}",
            f"ze spółką {CLIENT_B} zostało odnotowane w rejestrze banku.",
            "",
            "Z poważaniem",
            f"{PERSON_A}, {CLIENT_A}",
            "",
            DISCLAIMER,
        ],
        message_id="<demo-0002@finddocs.example>",
    )


def _build_skany(root: Path) -> None:
    """Tworzy katalog ze skanami bez warstwy tekstowej."""
    _write_scan_png(
        root / "skany" / "skan-potwierdzenia-wplaty.png",
        [
            "POTWIERDZENIE WPŁATY GOTÓWKOWEJ",
            "",
            f"{BANK}, oddział nr 00",
            "Data: 24.07.2015",
            f"Wpłacający: {PERSON_A}",
            "Kwota: 314 zł",
            "Rachunek: 99 8765 4321 0987",
            "                  6543 2109 8765",
            "Tytuł: opłata dodatkowa",
            "",
            "Podpis kasjera: .....................",
            "",
            DISCLAIMER,
        ],
    )
    _write_scan_pdf(
        root / "skany" / "skan-umowy-o-wspolpracy.pdf",
        [
            "UMOWA O WSPÓŁPRACY",
            "",
            "zawarta dnia 12.04.2005 pomiędzy",
            CLIENT_A,
            f"a {CLIENT_B}",
            "",
            "Przedmiotem umowy jest stała współpraca",
            "przy realizacji robót budowlanych.",
            "",
            "Rozliczenia prowadzone są w złotych,",
            "w terminie 14 dni od wystawienia faktury.",
            "",
            f"Podpisy stron: {PERSON_A}, {PERSON_D}",
            "",
            DISCLAIMER,
        ],
    )


def _build_raporty(root: Path, rng: random.Random) -> None:
    """Tworzy katalog z raportami."""
    rows: list[list[str]] = [
        ["Data raportu", "24.07.2015"],
        ["Rachunek", ACCOUNT_SPACED],
        ["Posiadacz", CLIENT_A],
        ["Liczba transakcji", "10"],
        ["Obroty uznaniowe", "24 750,00"],
        ["Obroty obciążeniowe", "21 643,50"],
        ["Największe obciążenie", "12 400,00"],
        ["Płatność kartą", "314,00"],
        ["Uwagi", DISCLAIMER],
    ]
    _write_xlsx(
        root / "raporty" / "raport-transakcji-2015.xlsx",
        [("Podsumowanie", ["pozycja", "wartość"], rows)],
        title="Raport transakcji lipiec 2015",
        author=PERSON_B,
    )
    _write_csv(
        root / "raporty" / "zestawienie-powiazan.csv",
        ["klient", "podmiot powiązany", "rodzaj powiązania", "rachunek", "data ustalenia"],
        [
            [
                CLIENT_A,
                CLIENT_B,
                f"poręczenie kredytu {CONTRACT_CREDIT}",
                ACCOUNT_SPACED,
                "2015-06-18",
            ],
            [
                CLIENT_B,
                CLIENT_A,
                "stała współpraca handlowa od 12.04.2005",
                ACCOUNT_B_SPACED,
                "2005-04-12",
            ],
            [
                CLIENT_A,
                "Biuro Serwis sp. z o.o.",
                "dostawca materiałów biurowych",
                ACCOUNT_TECH_SPACED,
                "2014-11-05",
            ],
            [DISCLAIMER, "", "", "", ""],
        ],
    )
    generated = _dt.datetime(2015, 8, 3, 7, 30)
    quarters = [
        ("I kwartał 2015", "184 320,00", "12"),
        ("II kwartał 2015", "201 480,50", "15"),
        ("III kwartał 2015", "97 640,25", str(9 + rng.randrange(0, 4))),
    ]
    html_rows = "\n".join(
        f"      <tr><td>{name}</td><td>{value} zł</td><td>{count}</td></tr>"
        for name, value, count in quarters
    )
    _write_text(
        root / "raporty" / "podsumowanie-kwartalne.html",
        [
            "<!DOCTYPE html>",
            '<html lang="pl">',
            "  <head>",
            '    <meta charset="utf-8" />',
            "    <title>Podsumowanie kwartalne 2015</title>",
            "  </head>",
            "  <body>",
            "    <h1>Podsumowanie kwartalne 2015</h1>",
            f"    <p>Raport wygenerowano {generated.strftime('%d.%m.%Y')} dla klienta "
            f"{CLIENT_A}.</p>",
            "    <table>",
            "      <tr><th>Okres</th><th>Obroty</th><th>Liczba transakcji</th></tr>",
            html_rows,
            "    </table>",
            f"    <p>Rachunek rozliczeniowy podmiotu powiązanego: {ACCOUNT_B_SPACED}.</p>",
            f"    <p>Sporna płatność kartą {CARD_MASKED} na kwotę 314 zł oczekuje na "
            "wyjaśnienie od 24 lipca 2015 r. (24.07.2015).</p>",
            f"    <p>Raport przygotowała {PERSON_B}.</p>",
            f"    <p>{DISCLAIMER}</p>",
            "  </body>",
            "</html>",
        ],
    )


def _build_problemy(root: Path, rng: random.Random) -> None:
    """Tworzy katalog z plikami celowo problematycznymi."""
    folder = root / "problemy"
    folder.mkdir(parents=True, exist_ok=True)
    broken = bytearray(b"%PDF-1.7\n%\xe2\xe3\xcf\xd3\n")
    broken += bytes(rng.randrange(256) for _ in range(4096))
    (folder / "uszkodzony-raport.pdf").write_bytes(bytes(broken))
    (folder / "pusty-dokument.txt").write_bytes(b"")
    # Wlasny format binarny nieznany aplikacji: bajty zerowe wykluczaja heurystyke tekstowa.
    export = bytearray(b"FDX\x01\x00\x02\x00\x00")
    for record in range(1, 25):
        export += struct.pack("<HHI", record, rng.randrange(1, 900), rng.randrange(1, 10**6))
        export += b"\x00\x1f"
    export += b"\x00\x00FDXEND\x00"
    (folder / "dane-eksportu.xyz").write_bytes(bytes(export))
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:body><w:p><w:r><w:t>Umowa zabezpieczona hasłem. "
        "Dokument testowy, dane fikcyjne.</w:t></w:r></w:p></w:body></w:document>"
    )
    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" ContentType="application/vnd.'
        'openxmlformats-officedocument.wordprocessingml.document.main+xml"/></Types>'
    )
    relationships = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/'
        'relationships/officeDocument" Target="word/document.xml"/></Relationships>'
    )
    write_protected_zip(
        folder / "zabezpieczona-umowa.docx",
        [
            ("[Content_Types].xml", content_types.encode("utf-8")),
            ("_rels/.rels", relationships.encode("utf-8")),
            ("word/document.xml", document_xml.encode("utf-8")),
        ],
        phrase="demo-2015",
        rng=rng,
        moment=_dt.datetime(2015, 7, 24, 9, 30),
    )


# --- zapytania referencyjne -----------------------------------------------------


def _query(query: str, mode: str, paths: tuple[str, ...], description: str) -> DemoQuery:
    """Skrot budujacy opis zapytania referencyjnego."""
    return DemoQuery(query=query, mode=mode, expected_paths=list(paths), description=description)


def reference_queries() -> list[DemoQuery]:
    """Zapytania referencyjne pokrywajace trzy tryby wyszukiwania."""
    return [
        _query(
            ACCOUNT_SPACED,
            "exact",
            ACCOUNT_DOCUMENTS_SPACED,
            "Numer rachunku zapisany ze spacjami, wariant najczestszy w zbiorze.",
        ),
        _query(
            ACCOUNT_DASHED,
            "exact",
            ACCOUNT_DOCUMENTS_DASHED,
            "Ten sam numer rachunku zapisany z myslnikami.",
        ),
        _query(
            ACCOUNT_COMPACT,
            "exact",
            ACCOUNT_DOCUMENTS_COMPACT,
            "Ten sam numer rachunku zapisany ciagiem cyfr.",
        ),
        _query(
            "24.07.2015",
            "exact",
            (
                DOC_PROC_PRZELEWY,
                DOC_POROZUMIENIE,
                DOC_MAIL_POTWIERDZENIE,
                DOC_RAPORT_XLSX,
                DOC_RAPORT_HTML,
            ),
            "Data w formacie dziennym, uzywana w procedurze przelewow.",
        ),
        _query(
            "2015-07-24",
            "exact",
            (DOC_PROC_PRZELEWY, DOC_TRX_A_2015, DOC_TRX_B_2015, DOC_MAIL_POTWIERDZENIE),
            "Ta sama data w formacie ISO, glownie w plikach transakcyjnych.",
        ),
        _query(
            "24 lipca 2015",
            "exact",
            (DOC_PROC_PRZELEWY, DOC_POROZUMIENIE, DOC_MAIL_POTWIERDZENIE, DOC_RAPORT_HTML),
            "Ta sama data zapisana slownie.",
        ),
        _query(
            "05.05.2007",
            "exact",
            (
                DOC_PROC_KSIEGOWANIE,
                DOC_ANEKS,
                DOC_TRX_A_2007,
                DOC_MAIL_ZAPYTANIE,
                DOC_UMOWA_RACHUNKU,
            ),
            "Data zdarzenia z 2007 roku, kluczowa dla zapytania o przebieg dnia.",
        ),
        _query(
            "2007-05-05",
            "exact",
            (DOC_PROC_KSIEGOWANIE, DOC_MAIL_ZAPYTANIE),
            "Ta sama data w formacie ISO.",
        ),
        _query(
            "5 maja 2007",
            "exact",
            (DOC_PROC_KSIEGOWANIE, DOC_ANEKS, DOC_MAIL_ZAPYTANIE),
            "Ta sama data zapisana slownie.",
        ),
        _query(
            PERSON_A,
            "exact",
            (DOC_PROC_PRZELEWY, DOC_UMOWA_RACHUNKU, DOC_ANEKS, DOC_MAIL_ZAPYTANIE),
            "Nazwisko i imie osoby wystepujacej w kilku formatach plikow.",
        ),
        _query(
            PERSON_C,
            "exact",
            (DOC_PROC_PRZELEWY, DOC_PROC_REKLAMACJE, DOC_UMOWA_KREDYTU),
            "Nazwisko z polskimi znakami diakrytycznymi.",
        ),
        _query(
            CLIENT_A,
            "exact",
            (
                DOC_PROC_AML,
                DOC_UMOWA_RACHUNKU,
                DOC_UMOWA_KREDYTU,
                DOC_ANEKS,
                DOC_POROZUMIENIE,
                DOC_TRX_A_2007,
                DOC_TRX_B_2015,
                DOC_TRX_B_XLSX,
                DOC_MAIL_POTWIERDZENIE,
                DOC_MAIL_ZAPYTANIE,
                DOC_RAPORT_CSV,
                DOC_RAPORT_HTML,
            ),
            "Pelna nazwa klienta A jako fraza z kropkami i skrotami.",
        ),
        _query(
            "Nowak-Bud",
            "exact",
            (
                DOC_PROC_AML,
                DOC_UMOWA_KREDYTU,
                DOC_POROZUMIENIE,
                DOC_TRX_A_2015,
                DOC_TRX_B_XLSX,
                DOC_MAIL_ZAPYTANIE,
                DOC_RAPORT_CSV,
            ),
            "Nazwa klienta B z myslnikiem w srodku wyrazu.",
        ),
        _query(
            "314 zł",
            "exact",
            (DOC_PROC_REKLAMACJE, DOC_POROZUMIENIE, DOC_RAPORT_HTML),
            "Kwota zapisana slowem zl, wymaga dopasowania jednostki.",
        ),
        _query(
            "314,00",
            "exact",
            (
                DOC_PROC_REKLAMACJE,
                DOC_TRX_A_2015,
                DOC_TRX_A_2007,
                DOC_MAIL_ZAPYTANIE,
                DOC_RAPORT_XLSX,
            ),
            "Ta sama kwota w zapisie ksiegowym z przecinkiem dziesietnym.",
        ),
        _query(
            CARD_SUFFIX,
            "exact",
            (DOC_PROC_REKLAMACJE, DOC_TRX_A_2015, DOC_RAPORT_HTML),
            "Koncowka numeru karty, fragment ciagu cyfr wewnatrz maski.",
        ),
        _query(
            "płatność kartą",
            "exact",
            (DOC_PROC_REKLAMACJE, DOC_TRX_A_2015, DOC_RAPORT_XLSX, DOC_RAPORT_HTML),
            "Fraza z polskimi znakami, w dokumentach wystepuje z wielkiej litery.",
        ),
        _query(
            "polecenie przelewu",
            "exact",
            (DOC_PROC_PRZELEWY, DOC_TRX_A_2015),
            "Fraza dwuwyrazowa z terminologii bankowej.",
        ),
        _query(
            "sesja rozliczeniowa",
            "exact",
            (DOC_PROC_PRZELEWY, DOC_PROC_KSIEGOWANIE),
            "Fraza w mianowniku, w tekstach wystepuje takze w odmianie.",
        ),
        _query(
            INVOICE_ID,
            "exact",
            (DOC_POROZUMIENIE, DOC_TRX_A_2015, DOC_TRX_B_2015, DOC_MAIL_POTWIERDZENIE),
            "Identyfikator faktury z ukosnikami.",
        ),
        _query(
            CONTRACT_CREDIT,
            "exact",
            (DOC_UMOWA_KREDYTU, DOC_POROZUMIENIE, DOC_TRX_B_2015, DOC_TRX_B_XLSX, DOC_RAPORT_CSV),
            "Numer umowy kredytowej z ukosnikami.",
        ),
        _query(
            CONTRACT_ACCOUNT,
            "exact",
            (DOC_UMOWA_RACHUNKU, DOC_ANEKS),
            "Numer umowy rachunku, wystepuje takze w tabeli dokumentu DOCX.",
        ),
        _query(
            DISCLAIMER,
            "exact",
            DISCLAIMER_DOCUMENTS,
            "Zdanie powtorzone w kazdym dokumencie, test kompletnosci indeksu.",
        ),
        _query(
            "Jak przekazać pieniądze innej firmie?",
            "semantic",
            (DOC_PROC_PRZELEWY,),
            "Parafraza bez wspolnych slow z tytulem procedury przelewow.",
        ),
        _query(
            "Co zrobić, gdy klient kwestionuje obciążenie konta?",
            "semantic",
            (DOC_PROC_REKLAMACJE,),
            "Parafraza opisu reklamacji, bez slowa reklamacja.",
        ),
        _query(
            "Zasady przechowywania i niszczenia starych akt",
            "semantic",
            (DOC_PROC_ARCHIWIZACJA,),
            "Parafraza procedury archiwizacji.",
        ),
        _query(
            "Kto ręczy za zobowiązania spółki budowlanej?",
            "semantic",
            (DOC_UMOWA_KREDYTU, DOC_PROC_AML),
            "Parafraza poreczenia kredytu, bez slowa poreczyciel.",
        ),
        _query(
            "Sprawdzanie tożsamości nowego kontrahenta i pochodzenia majątku",
            "semantic",
            (DOC_PROC_AML,),
            "Parafraza procedury poznania klienta.",
        ),
        _query(
            "Usterka oprogramowania i ręczne wprowadzanie zapisów",
            "semantic",
            (DOC_PROC_KSIEGOWANIE,),
            "Parafraza opisu awarii systemu ksiegowego.",
        ),
        _query(
            "Zakupy opłacone plastikiem w hurtowni",
            "semantic",
            (DOC_PROC_REKLAMACJE, DOC_TRX_A_2015),
            "Parafraza platnosci kartowej bez slowa karta.",
        ),
        _query(
            "Wspólne rozliczenia dwóch przedsiębiorstw",
            "semantic",
            (DOC_POROZUMIENIE, DOC_RAPORT_CSV),
            "Parafraza porozumienia rozliczeniowego.",
        ),
        _query(
            "Wiadomość z banku o wykonanej dyspozycji",
            "semantic",
            (DOC_MAIL_POTWIERDZENIE,),
            "Parafraza potwierdzenia przelewu wraz z zalacznikiem.",
        ),
        _query(
            "Ile kosztuje utrzymanie konta firmowego?",
            "semantic",
            (DOC_UMOWA_RACHUNKU, DOC_ANEKS),
            "Parafraza oplaty za prowadzenie rachunku.",
        ),
        _query(
            "Jaka była procedura dotycząca przelewów w dniu 24.07.2015?",
            "hybrid",
            (DOC_PROC_PRZELEWY, DOC_POROZUMIENIE),
            "Zapytanie ze specyfikacji: tresc merytoryczna wraz z data.",
        ),
        _query(
            f"Wyszukaj wszystkie transakcje z rachunku {ACCOUNT_SPACED}",
            "hybrid",
            (
                DOC_TRX_A_2015,
                DOC_TRX_B_2015,
                DOC_RAPORT_XLSX,
                DOC_RAPORT_CSV,
                DOC_MAIL_POTWIERDZENIE,
            ),
            "Zapytanie ze specyfikacji: numer rachunku w zdaniu naturalnym.",
        ),
        _query(
            "Co się działo w dniu 05.05.2007?",
            "hybrid",
            (DOC_PROC_KSIEGOWANIE, DOC_TRX_A_2007, DOC_ANEKS, DOC_MAIL_ZAPYTANIE),
            "Zapytanie ze specyfikacji: pytanie o przebieg konkretnego dnia.",
        ),
        _query(
            "Czy klient A miał powiązania z klientem B?",
            "hybrid",
            (DOC_RAPORT_CSV, DOC_UMOWA_KREDYTU, DOC_PROC_AML, DOC_TRX_B_XLSX, DOC_POROZUMIENIE),
            "Zapytanie ze specyfikacji: powiazania miedzy podmiotami.",
        ),
        _query(
            f"płatność kartą ...{CARD_SUFFIX}",
            "hybrid",
            (DOC_PROC_REKLAMACJE, DOC_TRX_A_2015, DOC_RAPORT_HTML),
            "Zapytanie ze specyfikacji: fragment numeru karty z opisem slownym.",
        ),
        _query(
            f"kwota 314 zł na rachunku {ACCOUNT_SPACED}",
            "hybrid",
            (DOC_PROC_REKLAMACJE, DOC_TRX_A_2015),
            "Kwota i numer rachunku w jednym zapytaniu.",
        ),
        _query(
            "umowa kredytu z poręczeniem spółki ACME z 2015 roku",
            "hybrid",
            (DOC_UMOWA_KREDYTU, DOC_RAPORT_CSV),
            "Tresc umowy wraz z rokiem i nazwa podmiotu.",
        ),
        _query(
            "skan potwierdzenia wpłaty na kwotę 314 zł",
            "hybrid",
            (DOC_SKAN_PNG,),
            "Wymaga OCR: tekst istnieje wylacznie na obrazie PNG.",
        ),
        _query(
            "zeskanowana umowa o współpracy z firmą budowlaną",
            "hybrid",
            (DOC_SKAN_PDF,),
            "Wymaga OCR: PDF bez warstwy tekstowej.",
        ),
    ]


# --- generowanie zbioru ---------------------------------------------------------


def generate_demo_corpus(
    target: Path,
    *,
    include_scans: bool = True,
    seed: int = 20240101,
) -> DemoCorpusInfo:
    """Tworzy zbior demonstracyjny w katalogu ``target`` i zwraca jego opis.

    Tresc dokumentow jest powtarzalna dla tego samego ziarna ``seed``. Same pliki
    nie musza byc identyczne bajt po bajcie, bo pakiety DOCX i XLSX zapisuja
    w archiwum czas zapisu. Gdy ``include_scans`` jest wylaczone, katalog ``skany``
    nie powstaje, dzieki czemu mozna testowac aplikacje bez silnika OCR.
    """
    root = Path(target).expanduser()
    if root.exists() and not root.is_dir():
        raise ConfigurationError(
            f"Sciezka {root} istnieje i nie jest katalogiem, nie mozna zapisac zbioru "
            "demonstracyjnego."
        )
    rng = random.Random(seed)
    try:
        root.mkdir(parents=True, exist_ok=True)
        _build_procedury(root)
        _build_umowy(root)
        _build_transakcje(root, rng)
        _build_korespondencja(root)
        if include_scans:
            _build_skany(root)
        _build_raporty(root, rng)
        _build_problemy(root, rng)
    except OSError as exc:
        raise TemporaryStorageError(
            f"Nie udalo sie zapisac zbioru demonstracyjnego w katalogu {root}.",
            details={"katalog": str(root)},
            cause=exc,
        ) from exc

    files = [
        path for path in sorted(root.rglob("*")) if path.is_file() and path.name != MANIFEST_NAME
    ]
    by_extension: dict[str, int] = {}
    for path in files:
        extension = path.suffix.lower()
        by_extension[extension] = by_extension.get(extension, 0) + 1

    info = DemoCorpusInfo(
        root=root,
        files=len(files),
        by_extension=dict(sorted(by_extension.items())),
        account_number=ACCOUNT_SPACED,
        account_documents=EXPECTED_ACCOUNT_DOCUMENTS,
        broken_files=3,
        unsupported_files=1,
        scan_files=2 if include_scans else 0,
        expected_queries=reference_queries(),
    )
    log.info(
        "demo.corpus_generated",
        katalog=str(root),
        pliki=info.files,
        skany=info.scan_files,
        ziarno=seed,
    )
    return info


def demo_corpus_path(base: Path | None = None) -> Path:
    """Domyslna sciezka zbioru demonstracyjnego wewnatrz katalogu danych uzytkownika."""
    paths = AppPaths.default() if base is None else AppPaths.at(base)
    return paths.root / "demo"


def ensure_demo_corpus(base: Path | None = None, force: bool = False) -> DemoCorpusInfo:
    """Zwraca opis zbioru demonstracyjnego, generujac go tylko gdy to konieczne.

    Gdy katalog istnieje i zawiera czytelny manifest, zbior nie jest generowany
    ponownie. Parametr ``force`` usuwa istniejacy katalog i tworzy zbior od nowa.
    """
    root = demo_corpus_path(base)
    if root.exists() and force:
        shutil.rmtree(root, ignore_errors=True)
    if root.exists():
        existing = load_manifest(root)
        if existing is not None:
            log.info("demo.corpus_reused", katalog=str(root), pliki=existing.files)
            return existing
    info = generate_demo_corpus(root)
    save_manifest(info)
    return info


# --- manifest -------------------------------------------------------------------


def _query_to_dict(query: DemoQuery) -> dict[str, Any]:
    return {
        "zapytanie": query.query,
        "tryb": query.mode,
        "dokumenty": list(query.expected_paths),
        "opis": query.description,
    }


def _query_from_dict(data: dict[str, Any]) -> DemoQuery:
    documents = data.get("dokumenty", [])
    paths = [str(item) for item in documents] if isinstance(documents, list) else []
    return DemoQuery(
        query=str(data.get("zapytanie", "")),
        mode=str(data.get("tryb", "hybrid")),
        expected_paths=paths,
        description=str(data.get("opis", "")),
    )


def manifest_to_dict(info: DemoCorpusInfo) -> dict[str, Any]:
    """Zamienia opis zbioru na strukture zapisywana w manifescie."""
    return {
        "wersja_manifestu": MANIFEST_VERSION,
        "wersja_aplikacji": APP_VERSION,
        "opis": (
            "Fikcyjny zbior demonstracyjny FindDocs. Wszystkie nazwiska, firmy, numery "
            "i kwoty sa wymyslone na potrzeby testow."
        ),
        "katalog": str(info.root),
        "liczba_plikow": info.files,
        "pliki_wg_rozszerzenia": dict(info.by_extension),
        "numer_rachunku": info.account_number,
        "warianty_numeru_rachunku": list(ACCOUNT_VARIANTS),
        "dokumenty_z_numerem_rachunku": info.account_documents,
        "lista_dokumentow_z_numerem_rachunku": list(ACCOUNT_DOCUMENTS),
        "pliki_uszkodzone": info.broken_files,
        "pliki_nieobslugiwane": info.unsupported_files,
        "pliki_skanow": info.scan_files,
        "zapytania_referencyjne": [_query_to_dict(q) for q in info.expected_queries],
    }


def save_manifest(info: DemoCorpusInfo) -> Path:
    """Zapisuje manifest zbioru w jego katalogu glownym."""
    info.root.mkdir(parents=True, exist_ok=True)
    target = info.root / MANIFEST_NAME
    payload = json.dumps(manifest_to_dict(info), ensure_ascii=False, indent=2)
    try:
        target.write_text(payload + "\n", encoding="utf-8")
    except OSError as exc:
        raise TemporaryStorageError(
            f"Nie udalo sie zapisac manifestu zbioru demonstracyjnego w {target}.",
            cause=exc,
        ) from exc
    return target


def load_manifest(root: Path) -> DemoCorpusInfo | None:
    """Wczytuje manifest zbioru. Zwraca None, gdy manifestu nie ma albo jest nieczytelny."""
    target = Path(root) / MANIFEST_NAME
    if not target.exists():
        return None
    try:
        raw = json.loads(target.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        log.warning("demo.manifest_unreadable", katalog=str(root))
        return None
    if not isinstance(raw, dict):
        log.warning("demo.manifest_invalid", katalog=str(root))
        return None
    queries_raw = raw.get("zapytania_referencyjne", [])
    queries = (
        [_query_from_dict(item) for item in queries_raw if isinstance(item, dict)]
        if isinstance(queries_raw, list)
        else []
    )
    extensions_raw = raw.get("pliki_wg_rozszerzenia", {})
    extensions = (
        {str(k): int(v) for k, v in extensions_raw.items()}
        if isinstance(extensions_raw, dict)
        else {}
    )
    return DemoCorpusInfo(
        root=Path(root),
        files=int(raw.get("liczba_plikow", 0) or 0),
        by_extension=extensions,
        account_number=str(raw.get("numer_rachunku", ACCOUNT_SPACED)),
        account_documents=int(raw.get("dokumenty_z_numerem_rachunku", 0) or 0),
        broken_files=int(raw.get("pliki_uszkodzone", 0) or 0),
        unsupported_files=int(raw.get("pliki_nieobslugiwane", 0) or 0),
        scan_files=int(raw.get("pliki_skanow", 0) or 0),
        expected_queries=queries,
    )


__all__ = [
    "ACCOUNT_COMPACT",
    "ACCOUNT_DASHED",
    "ACCOUNT_DOCUMENTS",
    "ACCOUNT_SPACED",
    "ACCOUNT_VARIANTS",
    "DISCLAIMER",
    "EXPECTED_ACCOUNT_DOCUMENTS",
    "MANIFEST_NAME",
    "MANIFEST_VERSION",
    "DemoCorpusInfo",
    "DemoQuery",
    "build_image_pdf",
    "build_text_pdf",
    "demo_corpus_path",
    "ensure_demo_corpus",
    "generate_demo_corpus",
    "load_manifest",
    "manifest_to_dict",
    "reference_queries",
    "render_scan_image",
    "save_manifest",
    "write_protected_zip",
]
